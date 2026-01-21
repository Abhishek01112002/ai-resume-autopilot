import os
from typing import Dict, Any
from pathlib import Path
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ResumeGenerator:
    def __init__(self):
        self.output_dir = Path("uploads/generated_resumes")
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def generate_pdf(self, resume_data: Dict[str, Any], output_filename: str) -> str:
        """Generate PDF resume"""
        file_path = self.output_dir / f"{output_filename}.pdf"
        
        doc = SimpleDocTemplate(str(file_path), pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=12,
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=6,
            spaceBefore=12,
        )
        
        # Name and contact (if available)
        if resume_data.get("name"):
            story.append(Paragraph(resume_data["name"], title_style))
            story.append(Spacer(1, 0.1*inch))
        
        # Summary
        if resume_data.get("summary"):
            story.append(Paragraph("Summary", heading_style))
            story.append(Paragraph(resume_data["summary"], styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
        
        # Skills
        if resume_data.get("skills"):
            story.append(Paragraph("Skills", heading_style))
            skills_text = ", ".join(resume_data["skills"])
            story.append(Paragraph(skills_text, styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
        
        # Experience
        if resume_data.get("experience"):
            story.append(Paragraph("Experience", heading_style))
            for exp in resume_data["experience"]:
                exp_text = f"<b>{exp.get('role', 'N/A')}</b>"
                if exp.get("company"):
                    exp_text += f" - {exp.get('company')}"
                if exp.get("duration"):
                    exp_text += f" ({exp.get('duration')})"
                story.append(Paragraph(exp_text, styles['Normal']))
                if exp.get("description"):
                    story.append(Paragraph(exp.get("description"), styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
            story.append(Spacer(1, 0.1*inch))
        
        # Projects
        if resume_data.get("projects"):
            story.append(Paragraph("Projects", heading_style))
            for project in resume_data["projects"]:
                proj_text = f"<b>{project.get('name', 'Project')}</b>"
                if project.get("tech"):
                    proj_text += f" - {', '.join(project.get('tech', []))}"
                story.append(Paragraph(proj_text, styles['Normal']))
                if project.get("description"):
                    story.append(Paragraph(project.get("description"), styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
            story.append(Spacer(1, 0.1*inch))
        
        # Education
        if resume_data.get("education"):
            edu = resume_data["education"]
            story.append(Paragraph("Education", heading_style))
            edu_text = ""
            if edu.get("degree"):
                edu_text += f"<b>{edu.get('degree')}</b>"
            if edu.get("field"):
                edu_text += f" in {edu.get('field')}"
            if edu.get("institution"):
                edu_text += f" - {edu.get('institution')}"
            if edu.get("year"):
                edu_text += f" ({edu.get('year')})"
            if edu_text:
                story.append(Paragraph(edu_text, styles['Normal']))
        
        doc.build(story)
        return str(file_path)
    
    def generate_docx(self, resume_data: Dict[str, Any], output_filename: str) -> str:
        """Generate DOCX resume"""
        file_path = self.output_dir / f"{output_filename}.docx"
        
        doc = Document()
        
        # Name
        if resume_data.get("name"):
            name_para = doc.add_paragraph(resume_data["name"])
            name_para.runs[0].font.size = Pt(18)
            name_para.runs[0].bold = True
            name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        doc.add_paragraph()  # Spacing
        
        # Summary
        if resume_data.get("summary"):
            doc.add_heading("Summary", level=2)
            doc.add_paragraph(resume_data["summary"])
            doc.add_paragraph()  # Spacing
        
        # Skills
        if resume_data.get("skills"):
            doc.add_heading("Skills", level=2)
            skills_text = ", ".join(resume_data["skills"])
            doc.add_paragraph(skills_text)
            doc.add_paragraph()  # Spacing
        
        # Experience
        if resume_data.get("experience"):
            doc.add_heading("Experience", level=2)
            for exp in resume_data["experience"]:
                exp_para = doc.add_paragraph()
                exp_run = exp_para.add_run(exp.get('role', 'N/A'))
                exp_run.bold = True
                if exp.get("company"):
                    exp_para.add_run(f" - {exp.get('company')}")
                if exp.get("duration"):
                    exp_para.add_run(f" ({exp.get('duration')})")
                if exp.get("description"):
                    doc.add_paragraph(exp.get("description"), style='List Bullet')
            doc.add_paragraph()  # Spacing
        
        # Projects
        if resume_data.get("projects"):
            doc.add_heading("Projects", level=2)
            for project in resume_data["projects"]:
                proj_para = doc.add_paragraph()
                proj_run = proj_para.add_run(project.get('name', 'Project'))
                proj_run.bold = True
                if project.get("tech"):
                    proj_para.add_run(f" - {', '.join(project.get('tech', []))}")
                if project.get("description"):
                    doc.add_paragraph(project.get("description"), style='List Bullet')
            doc.add_paragraph()  # Spacing
        
        # Education
        if resume_data.get("education"):
            edu = resume_data["education"]
            doc.add_heading("Education", level=2)
            edu_para = doc.add_paragraph()
            if edu.get("degree"):
                edu_run = edu_para.add_run(edu.get("degree"))
                edu_run.bold = True
            if edu.get("field"):
                edu_para.add_run(f" in {edu.get('field')}")
            if edu.get("institution"):
                edu_para.add_run(f" - {edu.get('institution')}")
            if edu.get("year"):
                edu_para.add_run(f" ({edu.get('year')})")
        
        doc.save(str(file_path))
        return str(file_path)

